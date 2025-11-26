"""Report generation for PDF and DOCX formats."""
from pathlib import Path
from datetime import datetime

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ReportGenerator:
    """Generates PDF and DOCX reports from markdown-like content."""

    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_pdf(self, title: str, content: str, filename: str) -> Path:
        """Generate a PDF report."""
        filepath = self.output_dir / f"{filename}.pdf"

        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#1a365d')
        ))
        styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#2c5282')
        ))
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceBefore=6,
            spaceAfter=6,
            leading=14
        ))

        story = []

        # Title
        story.append(Paragraph(title, styles['CustomTitle']))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 20))

        # Parse content (simple markdown-like parsing)
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['CustomHeading']))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('- '):
                story.append(Paragraph(f"• {line[2:]}", styles['CustomBody']))
            elif line.startswith('* '):
                story.append(Paragraph(f"• {line[2:]}", styles['CustomBody']))
            else:
                story.append(Paragraph(line, styles['CustomBody']))

        doc.build(story)
        return filepath

    def generate_docx(self, title: str, content: str, filename: str) -> Path:
        """Generate a DOCX report."""
        filepath = self.output_dir / f"{filename}.docx"

        doc = Document()

        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.runs[0]
        meta_run.font.size = Pt(10)
        meta_run.font.italic = True

        doc.add_paragraph()  # Spacer

        # Parse content
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        doc.save(str(filepath))
        return filepath

    def generate(self, title: str, content: str, format: str, filename: str) -> Path:
        """Generate a report in the specified format."""
        if format.lower() == 'pdf':
            return self.generate_pdf(title, content, filename)
        elif format.lower() == 'docx':
            return self.generate_docx(title, content, filename)
        else:
            raise ValueError(f"Unsupported format: {format}. Use 'pdf' or 'docx'.")
